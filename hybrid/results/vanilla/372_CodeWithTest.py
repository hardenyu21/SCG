import re
import glob
from docx import Document

def task_func(directory_path: str) -> int:
    # Find all .docx files in the provided directory
    docx_files = glob.glob(f"{directory_path}/*.docx")
    files_processed = 0

    for file_path in docx_files:
        # Open the document
        doc = Document(file_path)
        modified = False

        # Iterate through each paragraph in the document
        for para in doc.paragraphs:
            # Find all double quotes in the paragraph text
            if '"' in para.text:
                # Replace each double quote with a backslash-escaped double quote
                new_text = para.text.replace('"', '\\"')
                para.text = new_text
                modified = True

        # If any changes were made, save the document
        if modified:
            doc.save(file_path)
            files_processed += 1

    return files_processed
import unittest
import shutil
import os
import doctest
import tempfile
class TestCases(unittest.TestCase):
    def setUp(self):
        self.base_tmp_dir = tempfile.mkdtemp()
        self.test_directory = f"{self.base_tmp_dir}/test/"
        if not os.path.exists(self.test_directory):
            os.makedirs(self.test_directory)
        test_data = {
            "file_1.docx": "This is a sample text without any double quotes.",
            "file_2.docx": "This is a \"sample\" text with double quotes.",
            "file_3.docx": r'This is a \"sample\" text with double quotes already protected.',
            "file_4.docx": "Hello \"world\"! How are you \"today\"?",
            "file_5.docx": "Testing \"multiple\" paragraphs.\n\nAnother paragraph with \"quotes\"."
        }
        # Create .docx files for each scenario
        for file_name, content in test_data.items():
            doc = Document()
            for paragraph in content.split("\n"):
                doc.add_paragraph(paragraph)
            doc.save(self.test_directory + file_name)
    def tearDown(self):
        if os.path.exists(self.test_directory):
            shutil.rmtree(self.test_directory)
    def read_docx_content(self, file_path):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def test_case_1(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_1.docx")
        self.assertEqual(content, "This is a sample text without any double quotes.")
    def test_case_2(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_2.docx")
        self.assertEqual(content, r'This is a \"sample\" text with double quotes.')
    def test_case_3(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_3.docx")
        self.assertEqual(content, r'This is a \"sample\" text with double quotes already protected.')
    def test_case_4(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_4.docx")
        self.assertEqual(content, r'Hello \"world\"! How are you \"today\"?')
    def test_case_5(self):
        result = task_func(self.test_directory)
        self.assertEqual(result, 5)
        content = self.read_docx_content(self.test_directory + "file_5.docx")
        self.assertEqual(content, 'Testing \\"multiple\\" paragraphs.\n\nAnother paragraph with \\"quotes\\".')


if __name__ == '__main__':
    unittest.main(argv=['first-arg-is-ignored'], exit=False)